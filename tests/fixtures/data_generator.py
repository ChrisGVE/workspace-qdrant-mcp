"""
Comprehensive test data generation framework for workspace-qdrant-mcp testing.

This module provides generators for creating realistic, varied, and edge-case test data
for all test suites. Includes synthetic documents, code files, binary documents, project
structures, and more.

Usage Examples:
    >>> from tests.fixtures.data_generator import (
    ...     SyntheticDocumentGenerator,
    ...     CodeFileGenerator,
    ...     ProjectStructureGenerator
    ... )

    >>> # Generate synthetic text documents
    >>> doc_gen = SyntheticDocumentGenerator(seed=42)
    >>> docs = doc_gen.generate_documents(count=10, size="medium")

    >>> # Generate realistic code files
    >>> code_gen = CodeFileGenerator(seed=42)
    >>> python_files = code_gen.generate_python_module(
    ...     num_classes=3,
    ...     num_functions=5
    ... )

    >>> # Generate a complete project structure
    >>> project_gen = ProjectStructureGenerator(seed=42)
    >>> with project_gen.create_project_tree(scale="medium") as project_path:
    ...     # Use project_path for testing
    ...     pass
"""

import hashlib
import json
import random
import string
import tempfile
from contextlib import contextmanager
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any, Literal

# Optional dependencies - gracefully handle missing packages
try:
    from faker import Faker
    FAKER_AVAILABLE = True
except ImportError:
    FAKER_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


# Size scales for test data generation
SIZE_SCALES: dict[str, dict[str, int]] = {
    "tiny": {"lines": 10, "chars": 500, "files": 5},
    "small": {"lines": 50, "chars": 2_000, "files": 20},
    "medium": {"lines": 200, "chars": 10_000, "files": 100},
    "large": {"lines": 1000, "chars": 50_000, "files": 500},
    "xlarge": {"lines": 5000, "chars": 250_000, "files": 2000},
}


@dataclass
class GeneratedDocument:
    """Container for generated document data."""

    content: str | bytes
    filename: str
    file_type: str
    metadata: dict[str, Any] = field(default_factory=dict)
    size_bytes: int = 0

    def __post_init__(self):
        """Calculate size after initialization."""
        if isinstance(self.content, str):
            self.size_bytes = len(self.content.encode("utf-8"))
        else:
            self.size_bytes = len(self.content)


@dataclass
class GeneratedCodeFile:
    """Container for generated code file data."""

    content: str
    filename: str
    language: str
    symbols: list[str] = field(default_factory=list)  # Function/class names
    imports: list[str] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


class BaseGenerator:
    """Base class for all data generators with reproducible randomization."""

    def __init__(self, seed: int | None = None):
        """
        Initialize base generator.

        Args:
            seed: Random seed for reproducible generation. If None, uses system time.
        """
        self.seed = seed or int(datetime.now().timestamp())
        self.random = random.Random(self.seed)

        if FAKER_AVAILABLE:
            self.faker = Faker()
            Faker.seed(self.seed)
        else:
            self.faker = None

    def _random_string(self, length: int = 10, charset: str = string.ascii_letters) -> str:
        """Generate a random string of specified length."""
        return "".join(self.random.choices(charset, k=length))

    def _random_datetime(
        self,
        start_date: datetime | None = None,
        end_date: datetime | None = None
    ) -> datetime:
        """Generate a random datetime between start and end dates."""
        if start_date is None:
            start_date = datetime.now() - timedelta(days=365)
        if end_date is None:
            end_date = datetime.now()

        time_between = end_date - start_date
        days_between = time_between.days
        random_days = self.random.randrange(days_between)
        return start_date + timedelta(days=random_days)


class SyntheticDocumentGenerator(BaseGenerator):
    """
    Generate synthetic text documents with realistic content.

    Creates documents with faker-generated text, proper structure,
    and realistic metadata.
    """

    def generate_document(
        self,
        size: Literal["tiny", "small", "medium", "large", "xlarge"] = "medium",
        doc_type: Literal["article", "note", "email", "code_comment", "documentation"] = "article"
    ) -> GeneratedDocument:
        """
        Generate a single synthetic document.

        Args:
            size: Size scale of the document
            doc_type: Type of document to generate

        Returns:
            GeneratedDocument with content and metadata
        """
        if not FAKER_AVAILABLE:
            # Fallback to simple generation
            content = self._generate_simple_text(SIZE_SCALES[size]["chars"])
        else:
            content = self._generate_faker_content(size, doc_type)

        filename = f"{self._random_string(8)}_{doc_type}.txt"

        metadata = {
            "author": self.faker.name() if self.faker else "Test Author",
            "created_at": self._random_datetime().isoformat(),
            "doc_type": doc_type,
            "size_scale": size,
            "word_count": len(content.split()),
        }

        return GeneratedDocument(
            content=content,
            filename=filename,
            file_type="text/plain",
            metadata=metadata
        )

    def generate_documents(
        self,
        count: int = 10,
        size: Literal["tiny", "small", "medium", "large", "xlarge"] = "medium"
    ) -> list[GeneratedDocument]:
        """
        Generate multiple synthetic documents.

        Args:
            count: Number of documents to generate
            size: Size scale for all documents

        Returns:
            List of GeneratedDocument instances
        """
        doc_types = ["article", "note", "email", "code_comment", "documentation"]
        return [
            self.generate_document(size=size, doc_type=self.random.choice(doc_types))
            for _ in range(count)
        ]

    def _generate_faker_content(
        self,
        size: Literal["tiny", "small", "medium", "large", "xlarge"],
        doc_type: str
    ) -> str:
        """Generate content using Faker library."""
        target_chars = SIZE_SCALES[size]["chars"]

        if doc_type == "article":
            paragraphs = []
            current_length = 0
            while current_length < target_chars:
                para = self.faker.paragraph(nb_sentences=self.random.randint(3, 8))
                paragraphs.append(para)
                current_length += len(para)
            return "\n\n".join(paragraphs)

        elif doc_type == "note":
            # Short note with title
            title = self.faker.sentence(nb_words=6)
            content = self.faker.text(max_nb_chars=target_chars)
            return f"{title}\n\n{content}"

        elif doc_type == "email":
            # Email format
            sender = self.faker.email()
            recipient = self.faker.email()
            subject = self.faker.sentence(nb_words=8)
            body = self.faker.text(max_nb_chars=target_chars)
            return f"From: {sender}\nTo: {recipient}\nSubject: {subject}\n\n{body}"

        elif doc_type == "code_comment":
            # Multi-line code comment
            lines = []
            current_length = 0
            while current_length < target_chars:
                line = f"# {self.faker.sentence()}"
                lines.append(line)
                current_length += len(line)
            return "\n".join(lines)

        else:  # documentation
            # Structured documentation with headers
            sections = []
            for _ in range(self.random.randint(3, 6)):
                section_title = f"## {self.faker.sentence(nb_words=4)}"
                section_content = self.faker.paragraph(nb_sentences=5)
                sections.append(f"{section_title}\n\n{section_content}")
            return "\n\n".join(sections)

    def _generate_simple_text(self, target_chars: int) -> str:
        """Fallback text generation without faker."""
        words = ["test", "data", "content", "sample", "document", "text", "example"]
        result = []
        current_length = 0

        while current_length < target_chars:
            word = self.random.choice(words)
            result.append(word)
            current_length += len(word) + 1

        return " ".join(result)


class CodeFileGenerator(BaseGenerator):
    """
    Generate realistic code files in various languages.

    Supports Python, JavaScript, Rust, and other languages with
    proper syntax and structure.
    """

    def generate_python_module(
        self,
        num_classes: int = 3,
        num_functions: int = 5,
        include_docstrings: bool = True,
        complexity: Literal["simple", "moderate", "complex"] = "moderate"
    ) -> GeneratedCodeFile:
        """
        Generate a Python module with classes and functions.

        Args:
            num_classes: Number of classes to generate
            num_functions: Number of standalone functions
            include_docstrings: Whether to include docstrings
            complexity: Code complexity level

        Returns:
            GeneratedCodeFile with Python code
        """
        lines = []
        symbols = []
        imports = []

        # Standard imports
        import_lines = [
            "from typing import Any, Optional, List, Dict",
            "from dataclasses import dataclass, field",
            "import json",
            "from pathlib import Path",
        ]

        if complexity in ["moderate", "complex"]:
            import_lines.extend([
                "import asyncio",
                "from datetime import datetime",
            ])

        lines.extend(import_lines)
        imports.extend(import_lines)
        lines.append("\n")

        # Generate classes
        for i in range(num_classes):
            class_name = f"{self._random_string(8, string.ascii_uppercase[:10]).capitalize()}Handler"
            symbols.append(class_name)

            class_lines = self._generate_python_class(
                class_name,
                include_docstrings,
                complexity
            )
            lines.extend(class_lines)
            lines.append("\n")

        # Generate standalone functions
        for i in range(num_functions):
            func_name = f"process_{self._random_string(6, string.ascii_lowercase)}"
            symbols.append(func_name)

            func_lines = self._generate_python_function(
                func_name,
                include_docstrings,
                complexity
            )
            lines.extend(func_lines)
            lines.append("\n")

        content = "\n".join(lines)
        filename = f"test_module_{self._random_string(6)}.py"

        return GeneratedCodeFile(
            content=content,
            filename=filename,
            language="python",
            symbols=symbols,
            imports=imports,
            metadata={
                "num_classes": num_classes,
                "num_functions": num_functions,
                "complexity": complexity,
                "has_docstrings": include_docstrings,
            }
        )

    def generate_javascript_module(
        self,
        num_functions: int = 5,
        use_typescript: bool = False
    ) -> GeneratedCodeFile:
        """Generate a JavaScript/TypeScript module."""
        lines = []
        symbols = []
        imports = []

        # Imports
        import_lines = [
            "import { EventEmitter } from 'events';",
            "import path from 'path';",
        ]
        lines.extend(import_lines)
        imports.extend(import_lines)
        lines.append("\n")

        # Generate functions
        for i in range(num_functions):
            func_name = f"handle{self._random_string(6, string.ascii_uppercase[:10]).capitalize()}"
            symbols.append(func_name)

            if use_typescript:
                lines.append(f"export function {func_name}(data: any): Promise<void> {{")
            else:
                lines.append(f"export function {func_name}(data) {{")

            lines.append(f"  const result = processData(data);")
            lines.append(f"  return result;")
            lines.append("}\n")

        content = "\n".join(lines)
        ext = ".ts" if use_typescript else ".js"
        filename = f"test_module_{self._random_string(6)}{ext}"

        return GeneratedCodeFile(
            content=content,
            filename=filename,
            language="typescript" if use_typescript else "javascript",
            symbols=symbols,
            imports=imports
        )

    def generate_rust_module(
        self,
        num_structs: int = 2,
        num_functions: int = 3
    ) -> GeneratedCodeFile:
        """Generate a Rust module."""
        lines = []
        symbols = []

        # Use statements
        lines.append("use std::collections::HashMap;")
        lines.append("use std::fs::File;")
        lines.append("use std::io::{self, Read};")
        lines.append("\n")

        # Generate structs
        for i in range(num_structs):
            struct_name = f"{self._random_string(6, string.ascii_uppercase[:10]).capitalize()}Data"
            symbols.append(struct_name)

            lines.append("#[derive(Debug, Clone)]")
            lines.append(f"pub struct {struct_name} {{")
            lines.append("    pub id: u64,")
            lines.append("    pub name: String,")
            lines.append("    pub data: Vec<u8>,")
            lines.append("}\n")

        # Generate functions
        for i in range(num_functions):
            func_name = f"process_{self._random_string(6, string.ascii_lowercase)}"
            symbols.append(func_name)

            lines.append(f"pub fn {func_name}(input: &str) -> Result<String, io::Error> {{")
            lines.append("    let result = input.to_uppercase();")
            lines.append("    Ok(result)")
            lines.append("}\n")

        content = "\n".join(lines)
        filename = f"test_module_{self._random_string(6)}.rs"

        return GeneratedCodeFile(
            content=content,
            filename=filename,
            language="rust",
            symbols=symbols,
            imports=["std::collections::HashMap", "std::fs::File", "std::io"]
        )

    def _generate_python_class(
        self,
        class_name: str,
        include_docstrings: bool,
        complexity: str
    ) -> list[str]:
        """Generate a Python class definition."""
        lines = []

        lines.append(f"class {class_name}:")
        if include_docstrings:
            lines.append('    """')
            if self.faker:
                lines.append(f"    {self.faker.sentence()}")
            else:
                lines.append(f"    Handler class for {class_name}.")
            lines.append('    """')

        # __init__ method
        lines.append("    def __init__(self, config: Dict[str, Any]):")
        if include_docstrings:
            lines.append('        """Initialize the handler."""')
        lines.append("        self.config = config")
        lines.append("        self.data: List[Any] = []")
        lines.append("")

        # Additional methods based on complexity
        num_methods = {"simple": 1, "moderate": 3, "complex": 5}[complexity]

        for i in range(num_methods):
            method_name = f"process_{self._random_string(6, string.ascii_lowercase)}"
            lines.append(f"    def {method_name}(self, input_data: Any) -> Optional[Dict[str, Any]]:")
            if include_docstrings:
                lines.append('        """Process input data."""')
            lines.append("        result = {'processed': True, 'data': input_data}")
            lines.append("        return result")
            lines.append("")

        return lines

    def _generate_python_function(
        self,
        func_name: str,
        include_docstrings: bool,
        complexity: str
    ) -> list[str]:
        """Generate a standalone Python function."""
        lines = []

        if complexity == "complex":
            lines.append("async def " + func_name + "(data: Dict[str, Any]) -> List[str]:")
        else:
            lines.append("def " + func_name + "(data: Dict[str, Any]) -> List[str]:")

        if include_docstrings:
            lines.append('    """')
            if self.faker:
                lines.append(f"    {self.faker.sentence()}")
            else:
                lines.append(f"    Process data and return results.")
            lines.append('    """')

        lines.append("    results = []")
        lines.append("    for key, value in data.items():")
        lines.append("        results.append(f'{key}: {value}')")
        lines.append("    return results")

        return lines


class BinaryDocumentGenerator(BaseGenerator):
    """
    Generate binary documents (PDF, DOCX).

    Requires optional dependencies: reportlab, python-docx
    """

    def generate_pdf(
        self,
        num_pages: int = 3,
        include_images: bool = False
    ) -> GeneratedDocument:
        """
        Generate a PDF document.

        Args:
            num_pages: Number of pages
            include_images: Whether to include images (not implemented)

        Returns:
            GeneratedDocument with PDF binary content
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF generation. Install with: pip install reportlab")

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_text = self.faker.sentence(nb_words=8) if self.faker else "Test Document"
        title = Paragraph(title_text, styles['Title'])
        story.append(title)
        story.append(Spacer(1, 0.2 * inch))

        # Content
        for page in range(num_pages):
            if self.faker:
                content_text = self.faker.text(max_nb_chars=1000)
            else:
                content_text = f"This is test content for page {page + 1}. " * 50

            para = Paragraph(content_text, styles['Normal'])
            story.append(para)
            story.append(Spacer(1, 0.3 * inch))

        doc.build(story)

        pdf_data = buffer.getvalue()
        buffer.close()

        filename = f"test_document_{self._random_string(6)}.pdf"

        return GeneratedDocument(
            content=pdf_data,
            filename=filename,
            file_type="application/pdf",
            metadata={
                "num_pages": num_pages,
                "created_at": datetime.now().isoformat(),
            }
        )

    def generate_docx(
        self,
        num_paragraphs: int = 10,
        include_tables: bool = False
    ) -> GeneratedDocument:
        """
        Generate a DOCX document.

        Args:
            num_paragraphs: Number of paragraphs
            include_tables: Whether to include tables

        Returns:
            GeneratedDocument with DOCX binary content
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")

        document = Document()

        # Title
        title_text = self.faker.sentence(nb_words=8) if self.faker else "Test Document"
        document.add_heading(title_text, 0)

        # Paragraphs
        for _ in range(num_paragraphs):
            if self.faker:
                para_text = self.faker.paragraph(nb_sentences=5)
            else:
                para_text = "This is a test paragraph. " * 10
            document.add_paragraph(para_text)

        # Optional table
        if include_tables:
            table = document.add_table(rows=3, cols=3)
            for row in table.rows:
                for cell in row.cells:
                    cell.text = self.faker.word() if self.faker else "Data"

        # Save to bytes
        buffer = BytesIO()
        document.save(buffer)
        docx_data = buffer.getvalue()
        buffer.close()

        filename = f"test_document_{self._random_string(6)}.docx"

        return GeneratedDocument(
            content=docx_data,
            filename=filename,
            file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata={
                "num_paragraphs": num_paragraphs,
                "has_tables": include_tables,
                "created_at": datetime.now().isoformat(),
            }
        )


class ProjectStructureGenerator(BaseGenerator):
    """
    Generate realistic project directory structures.

    Creates temporary project trees with various file types,
    realistic naming, and proper organization.
    """

    def __init__(self, seed: int | None = None):
        """Initialize project structure generator."""
        super().__init__(seed)
        self.doc_gen = SyntheticDocumentGenerator(seed=self.seed)
        self.code_gen = CodeFileGenerator(seed=self.seed)

    @contextmanager
    def create_project_tree(
        self,
        scale: Literal["tiny", "small", "medium", "large"] = "small",
        languages: list[str] | None = None,
        include_git: bool = True
    ):
        """
        Create a temporary project directory tree.

        Args:
            scale: Size scale of the project
            languages: List of languages to include (default: ["python"])
            include_git: Whether to include .git directory

        Yields:
            Path to the temporary project directory

        Example:
            >>> gen = ProjectStructureGenerator(seed=42)
            >>> with gen.create_project_tree(scale="small") as project_path:
            ...     # Use project_path for testing
            ...     assert (project_path / "src").exists()
        """
        if languages is None:
            languages = ["python"]

        with tempfile.TemporaryDirectory() as tmpdir:
            project_path = Path(tmpdir) / f"test_project_{self._random_string(6)}"
            project_path.mkdir()

            # Create directory structure
            self._create_directories(project_path, scale)

            # Populate with files
            self._populate_project(project_path, scale, languages)

            # Create git repository
            if include_git:
                self._create_git_structure(project_path)

            yield project_path

    def _create_directories(
        self,
        project_path: Path,
        scale: Literal["tiny", "small", "medium", "large"]
    ):
        """Create project directory structure."""
        # Standard directories
        (project_path / "src").mkdir()
        (project_path / "tests").mkdir()
        (project_path / "docs").mkdir()

        if scale in ["medium", "large"]:
            (project_path / "src" / "core").mkdir()
            (project_path / "src" / "utils").mkdir()
            (project_path / "tests" / "unit").mkdir()
            (project_path / "tests" / "integration").mkdir()

        if scale == "large":
            (project_path / "src" / "api").mkdir()
            (project_path / "src" / "models").mkdir()
            (project_path / "benchmarks").mkdir()
            (project_path / "examples").mkdir()

    def _populate_project(
        self,
        project_path: Path,
        scale: Literal["tiny", "small", "medium", "large"],
        languages: list[str]
    ):
        """Populate project with files."""
        num_files = SIZE_SCALES[scale]["files"]

        # Root files
        (project_path / "README.md").write_text("# Test Project\n\nThis is a test project.")
        (project_path / "LICENSE").write_text("MIT License\n\nCopyright (c) 2025")

        # Python files
        if "python" in languages:
            self._create_python_files(project_path, num_files // 2)

        # JavaScript files
        if "javascript" in languages:
            self._create_javascript_files(project_path, num_files // 4)

        # Documentation files
        self._create_documentation_files(project_path, min(5, num_files // 10))

        # Config files
        self._create_config_files(project_path)

    def _create_python_files(self, project_path: Path, count: int):
        """Create Python source files."""
        for i in range(count):
            code_file = self.code_gen.generate_python_module(
                num_classes=self.random.randint(1, 3),
                num_functions=self.random.randint(2, 5)
            )

            # Place in appropriate directory
            if i % 3 == 0:
                target_dir = project_path / "src"
            elif i % 3 == 1:
                target_dir = project_path / "src" / "core"
            else:
                target_dir = project_path / "tests" / "unit"

            if not target_dir.exists():
                target_dir = project_path / "src"

            (target_dir / code_file.filename).write_text(code_file.content)

    def _create_javascript_files(self, project_path: Path, count: int):
        """Create JavaScript source files."""
        for i in range(count):
            code_file = self.code_gen.generate_javascript_module(
                num_functions=self.random.randint(3, 6)
            )

            target_dir = project_path / "src"
            (target_dir / code_file.filename).write_text(code_file.content)

    def _create_documentation_files(self, project_path: Path, count: int):
        """Create documentation files."""
        for i in range(count):
            doc = self.doc_gen.generate_document(size="small", doc_type="documentation")
            (project_path / "docs" / doc.filename).write_text(doc.content)

    def _create_config_files(self, project_path: Path):
        """Create configuration files."""
        # pyproject.toml
        pyproject_content = """[build-system]
requires = ["setuptools>=61.0"]
build-backend = "setuptools.build_meta"

[project]
name = "test-project"
version = "0.1.0"
"""
        (project_path / "pyproject.toml").write_text(pyproject_content)

        # .gitignore
        gitignore_content = """__pycache__/
*.py[cod]
*.so
.env
venv/
"""
        (project_path / ".gitignore").write_text(gitignore_content)

    def _create_git_structure(self, project_path: Path):
        """Create minimal .git directory structure."""
        git_dir = project_path / ".git"
        git_dir.mkdir()

        # HEAD file
        (git_dir / "HEAD").write_text("ref: refs/heads/main\n")

        # Config
        config_content = """[core]
    repositoryformatversion = 0
    filemode = true
    bare = false
"""
        (git_dir / "config").write_text(config_content)


class MetadataGenerator(BaseGenerator):
    """
    Generate realistic metadata for files and documents.

    Includes Git metadata, file attributes, timestamps, and more.
    """

    def generate_git_metadata(self) -> dict[str, Any]:
        """Generate Git-related metadata."""
        branch_names = ["main", "develop", "feature/auth", "bugfix/issue-123", "release/v1.0"]

        return {
            "branch": self.random.choice(branch_names),
            "commit_hash": hashlib.sha1(
                self._random_string(20).encode()
            ).hexdigest()[:7],
            "author": self.faker.name() if self.faker else "Test Author",
            "author_email": self.faker.email() if self.faker else "test@example.com",
            "commit_date": self._random_datetime().isoformat(),
        }

    def generate_file_metadata(
        self,
        file_path: str | Path,
        file_size: int | None = None
    ) -> dict[str, Any]:
        """Generate file system metadata."""
        if file_size is None:
            file_size = self.random.randint(100, 1_000_000)

        return {
            "file_path": str(file_path),
            "file_name": Path(file_path).name,
            "file_size": file_size,
            "created_at": self._random_datetime().isoformat(),
            "modified_at": self._random_datetime().isoformat(),
            "file_type": self._detect_file_type(Path(file_path)),
        }

    def generate_project_metadata(
        self,
        project_name: str | None = None
    ) -> dict[str, Any]:
        """Generate project-level metadata."""
        if project_name is None:
            project_name = f"test-project-{self._random_string(6)}"

        return {
            "project_name": project_name,
            "project_id": hashlib.sha256(project_name.encode()).hexdigest()[:16],
            "language": self.random.choice(["python", "javascript", "rust", "go"]),
            "framework": self.random.choice(["fastapi", "flask", "express", "actix"]),
            "created_at": self._random_datetime().isoformat(),
            "last_modified": self._random_datetime().isoformat(),
        }

    def _detect_file_type(self, path: Path) -> str:
        """Detect file type from extension."""
        extension_map = {
            ".py": "text/x-python",
            ".js": "text/javascript",
            ".ts": "text/typescript",
            ".rs": "text/rust",
            ".md": "text/markdown",
            ".txt": "text/plain",
            ".json": "application/json",
            ".yaml": "application/yaml",
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        return extension_map.get(path.suffix, "application/octet-stream")


class EdgeCaseGenerator(BaseGenerator):
    """
    Generate edge cases for testing robustness.

    Includes empty files, very large files, special characters,
    unicode content, and other boundary conditions.
    """

    def generate_empty_file(self) -> GeneratedDocument:
        """Generate an empty file."""
        return GeneratedDocument(
            content="",
            filename=f"empty_{self._random_string(6)}.txt",
            file_type="text/plain",
            metadata={"edge_case": "empty_file"}
        )

    def generate_very_large_file(
        self,
        size_mb: int = 10
    ) -> GeneratedDocument:
        """Generate a very large file (in-memory, use with caution)."""
        target_size = size_mb * 1024 * 1024
        chunk = "x" * 1024  # 1KB chunks
        num_chunks = target_size // 1024

        content = chunk * num_chunks

        return GeneratedDocument(
            content=content,
            filename=f"large_{size_mb}mb_{self._random_string(6)}.txt",
            file_type="text/plain",
            metadata={"edge_case": "very_large", "size_mb": size_mb}
        )

    def generate_unicode_content(self) -> GeneratedDocument:
        """Generate file with unicode characters."""
        unicode_samples = [
            "Hello 世界 🌍",
            "Привет мир",
            "مرحبا بالعالم",
            "こんにちは世界",
            "שלום עולם",
            "Emoji: 🚀 🎉 💻 🔥 ✨",
        ]

        content = "\n\n".join(unicode_samples * 10)

        return GeneratedDocument(
            content=content,
            filename=f"unicode_{self._random_string(6)}.txt",
            file_type="text/plain",
            metadata={"edge_case": "unicode_content", "encoding": "utf-8"}
        )

    def generate_special_chars_filename(self) -> GeneratedDocument:
        """Generate file with special characters in filename."""
        special_chars = ["spaces in name", "under_scores", "dots.in.name", "UPPERCASE"]

        name_part = self.random.choice(special_chars)
        filename = f"{name_part}_{self._random_string(4)}.txt"

        return GeneratedDocument(
            content="Test content with special filename",
            filename=filename,
            file_type="text/plain",
            metadata={"edge_case": "special_chars_filename"}
        )

    def generate_binary_noise(self, size_bytes: int = 1024) -> GeneratedDocument:
        """Generate random binary data."""
        content = bytes(self.random.getrandbits(8) for _ in range(size_bytes))

        return GeneratedDocument(
            content=content,
            filename=f"binary_{self._random_string(6)}.bin",
            file_type="application/octet-stream",
            metadata={"edge_case": "binary_noise", "size_bytes": size_bytes}
        )


# Convenience function for quick data generation
def quick_generate(
    doc_type: Literal["text", "code", "binary", "project"] = "text",
    count: int = 10,
    seed: int | None = None,
    **kwargs
) -> list[GeneratedDocument] | list[GeneratedCodeFile] | Path:
    """
    Quick generation function for common use cases.

    Args:
        doc_type: Type of data to generate
        count: Number of items to generate
        seed: Random seed
        **kwargs: Additional arguments for specific generators
            For 'code': language, num_classes, num_functions, etc.
            For 'binary': binary_type, num_pages, num_paragraphs, etc.
            For 'text': size, doc_type

    Returns:
        List of generated items or project path

    Example:
        >>> docs = quick_generate("text", count=5, seed=42, size="small")
        >>> code_files = quick_generate("code", count=3, seed=42, language="python", num_classes=2)
        >>> binary_docs = quick_generate("binary", count=2, seed=42, binary_type="docx")
    """
    if doc_type == "text":
        gen = SyntheticDocumentGenerator(seed=seed)
        return gen.generate_documents(count=count, **kwargs)

    elif doc_type == "code":
        gen = CodeFileGenerator(seed=seed)
        language = kwargs.pop("language", "python")  # Remove language from kwargs

        if language == "python":
            return [gen.generate_python_module(**kwargs) for _ in range(count)]
        elif language in ["javascript", "typescript"]:
            use_typescript = language == "typescript"
            # Filter kwargs for javascript_module
            js_kwargs = {k: v for k, v in kwargs.items() if k in ["num_functions", "use_typescript"]}
            js_kwargs["use_typescript"] = use_typescript
            return [gen.generate_javascript_module(**js_kwargs) for _ in range(count)]
        elif language == "rust":
            # Filter kwargs for rust_module
            rust_kwargs = {k: v for k, v in kwargs.items() if k in ["num_structs", "num_functions"]}
            return [gen.generate_rust_module(**rust_kwargs) for _ in range(count)]
        else:
            raise ValueError(f"Unsupported language: {language}")

    elif doc_type == "binary":
        gen = BinaryDocumentGenerator(seed=seed)
        binary_type = kwargs.pop("binary_type", "pdf")  # Remove binary_type from kwargs

        if binary_type == "pdf":
            # Filter kwargs for pdf
            pdf_kwargs = {k: v for k, v in kwargs.items() if k in ["num_pages", "include_images"]}
            return [gen.generate_pdf(**pdf_kwargs) for _ in range(count)]
        elif binary_type == "docx":
            # Filter kwargs for docx
            docx_kwargs = {k: v for k, v in kwargs.items() if k in ["num_paragraphs", "include_tables"]}
            return [gen.generate_docx(**docx_kwargs) for _ in range(count)]
        else:
            raise ValueError(f"Unsupported binary type: {binary_type}")

    elif doc_type == "project":
        gen = ProjectStructureGenerator(seed=seed)
        # For project, we can't return a context manager directly
        # Users should use create_project_tree() directly
        raise ValueError(
            "Use ProjectStructureGenerator.create_project_tree() context manager directly for project generation"
        )

    else:
        raise ValueError(f"Unknown doc_type: {doc_type}")
