from ...observability import get_logger

logger = get_logger(__name__)
"""
Microsoft Word DOCX document parser for extracting text content and metadata.

This module provides functionality to parse DOCX Word documents, extracting
text content from paragraphs, tables, headers, and footers while preserving
document metadata and structure information.
"""

import logging
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .base import DocumentParser, ParsedDocument

logger = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    """
    Parser for Microsoft Word DOCX documents.

    Extracts text content from paragraphs, tables, headers, and footers
    while preserving document metadata like title, author, and creation date.
    """

    @property
    def supported_extensions(self) -> list[str]:
        """DOCX file extensions."""
        return [".docx"]

    @property
    def format_name(self) -> str:
        """Human-readable format name."""
        return "Microsoft Word DOCX"

    def _check_availability(self) -> None:
        """Check if required libraries are available."""
        if not DOCX_AVAILABLE:
            raise RuntimeError(
                "DOCX parsing requires 'python-docx'. "
                "Install with: pip install python-docx"
            )

    async def parse(self, file_path: str | Path, **options: Any) -> ParsedDocument:
        """
        Parse DOCX file and extract text content.

        Args:
            file_path: Path to DOCX file
            **options: Parsing options
                - include_tables: bool = True - Include table content
                - include_headers_footers: bool = True - Include headers/footers
                - table_format: str = "markdown" - Table formatting ("markdown", "plain", "csv")
                - preserve_formatting: bool = False - Preserve basic formatting markers

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            RuntimeError: If required libraries are not installed
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        self._check_availability()
        self.validate_file(file_path)

        file_path = Path(file_path)

        try:
            # Parse options
            include_tables = options.get("include_tables", True)
            include_headers_footers = options.get("include_headers_footers", True)
            table_format = options.get("table_format", "markdown")
            preserve_formatting = options.get("preserve_formatting", False)

            # Read DOCX document
            doc = Document(str(file_path))

            # Extract metadata
            metadata = await self._extract_metadata(doc)

            # Extract text content
            text_content = await self._extract_text_content(
                doc,
                include_tables,
                include_headers_footers,
                table_format,
                preserve_formatting,
            )

            # Parsing information
            parsing_info = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
                "content_length": len(text_content),
                "include_tables": include_tables,
                "include_headers_footers": include_headers_footers,
                "table_format": table_format,
            }

            logger.info(
                f"Successfully parsed DOCX: {file_path.name} "
                f"({parsing_info['paragraph_count']} paragraphs, "
                f"{parsing_info['table_count']} tables, "
                f"{parsing_info['content_length']:,} characters)"
            )

            return ParsedDocument.create(
                content=text_content,
                file_path=file_path,
                file_type="docx",
                additional_metadata=metadata,
                parsing_info=parsing_info,
            )

        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise RuntimeError(f"DOCX parsing failed: {e}") from e

    async def _extract_metadata(
        self, doc: Document
    ) -> dict[str, str | int | float | bool]:
        """Extract metadata from DOCX document."""
        metadata = {}

        # Core document properties
        core_props = doc.core_properties

        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.subject:
            metadata["subject"] = core_props.subject
        if core_props.keywords:
            metadata["keywords"] = core_props.keywords
        if core_props.comments:
            metadata["comments"] = core_props.comments
        if core_props.category:
            metadata["category"] = core_props.category
        if core_props.language:
            metadata["language"] = core_props.language

        # Dates
        if core_props.created:
            metadata["created_date"] = core_props.created.isoformat()
        if core_props.modified:
            metadata["modified_date"] = core_props.modified.isoformat()
        if core_props.last_modified_by:
            metadata["last_modified_by"] = core_props.last_modified_by

        # Document statistics
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)
        metadata["section_count"] = len(doc.sections)

        # Check for embedded content
        has_images = any(
            run._element.xpath(
                ".//a:blip",
                namespaces={
                    "a": "http://schemas.openxmlformats.org/drawingml/2006/main"
                },
            )
            for para in doc.paragraphs
            for run in para.runs
        )
        metadata["has_images"] = has_images

        # Check for hyperlinks
        has_hyperlinks = any(
            para._element.xpath(
                ".//w:hyperlink",
                namespaces={
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                },
            )
            for para in doc.paragraphs
        )
        metadata["has_hyperlinks"] = has_hyperlinks

        return metadata

    async def _extract_text_content(
        self,
        doc: Document,
        include_tables: bool,
        include_headers_footers: bool,
        table_format: str,
        preserve_formatting: bool,
    ) -> str:
        """Extract text content from the document."""
        content_parts = []

        # Extract headers and footers if requested
        if include_headers_footers:
            headers_footers = await self._extract_headers_footers(doc)
            if headers_footers:
                content_parts.append("=== HEADERS AND FOOTERS ===")
                content_parts.append(headers_footers)
                content_parts.append("=== DOCUMENT CONTENT ===")

        # Extract main document content
        main_content = await self._extract_main_content(
            doc, include_tables, table_format, preserve_formatting
        )
        if main_content:
            content_parts.append(main_content)

        return "\n\n".join(content_parts)

    async def _extract_main_content(
        self,
        doc: Document,
        include_tables: bool,
        table_format: str,
        preserve_formatting: bool,
    ) -> str:
        """Extract content from the main document body."""
        content_parts = []

        # Iterate through document elements in order
        for element in doc.element.body:
            if element.tag.endswith("p"):  # Paragraph
                paragraph = Paragraph(element, doc)
                text = await self._extract_paragraph_text(
                    paragraph, preserve_formatting
                )
                if text.strip():
                    content_parts.append(text)

            elif element.tag.endswith("tbl") and include_tables:  # Table
                table = Table(element, doc)
                table_text = await self._extract_table_text(table, table_format)
                if table_text:
                    content_parts.append(table_text)

        return "\n\n".join(content_parts)

    async def _extract_paragraph_text(
        self, paragraph: Paragraph, preserve_formatting: bool
    ) -> str:
        """Extract text from a paragraph with optional formatting preservation."""
        if not preserve_formatting:
            return paragraph.text

        # Preserve basic formatting with simple markers
        text_parts = []
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue

            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"_{text}_"

            text_parts.append(text)

        return "".join(text_parts)

    async def _extract_table_text(self, table: Table, table_format: str) -> str:
        """Extract text from a table with specified formatting."""
        if not table.rows:
            return ""

        rows_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Extract text from all paragraphs in the cell
                cell_text = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        cell_text.append(paragraph.text.strip())
                cell_content = " ".join(cell_text)
                row_data.append(cell_content)
            rows_data.append(row_data)

        if not rows_data:
            return ""

        if table_format == "markdown":
            return await self._format_table_markdown(rows_data)
        elif table_format == "csv":
            return await self._format_table_csv(rows_data)
        else:  # plain
            return await self._format_table_plain(rows_data)

    async def _format_table_markdown(self, rows_data: list[list[str]]) -> str:
        """Format table data as markdown."""
        if not rows_data:
            return ""

        lines = []

        # Header row
        if rows_data:
            header = " | ".join(rows_data[0])
            lines.append(f"| {header} |")

            # Separator
            separator = " | ".join("---" for _ in rows_data[0])
            lines.append(f"| {separator} |")

            # Data rows
            for row in rows_data[1:]:
                row_text = " | ".join(cell.replace("|", "\\|") for cell in row)
                lines.append(f"| {row_text} |")

        return "\n".join(lines)

    async def _format_table_csv(self, rows_data: list[list[str]]) -> str:
        """Format table data as CSV."""
        import csv
        from io import StringIO

        output = StringIO()
        writer = csv.writer(output)
        for row in rows_data:
            writer.writerow(row)
        return output.getvalue().strip()

    async def _format_table_plain(self, rows_data: list[list[str]]) -> str:
        """Format table data as plain text."""
        lines = []
        for row in rows_data:
            lines.append("\t".join(row))
        return "\n".join(lines)

    async def _extract_headers_footers(self, doc: Document) -> str:
        """Extract text from headers and footers."""
        content_parts = []

        for section in doc.sections:
            # Headers
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(f"Header: {paragraph.text.strip()}")

            # Footers
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(f"Footer: {paragraph.text.strip()}")

        return "\n".join(content_parts)

    def get_parsing_options(self) -> dict[str, dict[str, Any]]:
        """Get available parsing options for DOCX files."""
        return {
            "include_tables": {
                "type": bool,
                "default": True,
                "description": "Include table content in extracted text",
            },
            "include_headers_footers": {
                "type": bool,
                "default": True,
                "description": "Include headers and footers in extracted text",
            },
            "table_format": {
                "type": str,
                "default": "markdown",
                "description": "Table formatting style (markdown, plain, csv)",
            },
            "preserve_formatting": {
                "type": bool,
                "default": False,
                "description": "Preserve basic formatting with simple markers",
            },
        }
